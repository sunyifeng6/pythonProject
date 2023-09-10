
import os
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import Cm

title_name = ["一、","二、","三、","四、","五、","六、","七、","八、","九、","十","十一、","十二、","十三、"]


# 创建一个函数，将文档中所有的图片保存到images文件夹下，图片名以图片的embed属性值命名（即 rId 命名)
def extract_and_save_images_from_docx(doc_path):
    # 加载文档
    doc = Document(doc_path)

    # 创建用于存储图片的文件夹
    image_folder = 'images'
    if not os.path.exists(image_folder):
        os.makedirs(image_folder)

    # 获取文档中的所有图片
    def find_images(document):
        images = []
        rels = document.part.rels
        for rel in rels.values():
            if 'image' in rel.reltype:
                image_part = rel.target_part
                embed_id = rel.rId
                images.append((image_part, embed_id))
        return images

    # 保存图片
    image_list = find_images(doc)
    for i, (image_part, embed_id) in enumerate(image_list):
        image_filename = f'{embed_id}.png'
        image_path = os.path.join(image_folder, image_filename)
        with open(image_path, 'wb') as f:
            f.write(image_part.blob)

    print("Images extracted and saved successfully.")


# 创建分化文档的函数，新文件名以  标号+标题名命名 ，标号每次匹配到标题自动 +1
def extract_images_from_docx(doc_path):
    # 加载文档
    doc = Document(doc_path)

    # 创建新文档
    new_doc = Document()
    doc = Document(doc_path)
    current_title = None
    new_doc = None

    # 创建用于存储文档的文件夹
    file_folder = '零散-25350'
    if not os.path.exists(file_folder):
        os.makedirs(file_folder)



    # 遍历段落
    for paragraph in doc.paragraphs:
        if paragraph.text[:2] in title_name:  # 判断是否是标题
            if new_doc is not None:
                file_name = current_title + '.docx'
                file_path = os.path.join(file_folder, file_name)
                new_doc.save(file_path)
                # new_doc.save(f'{current_index}{current_title}.docx') # 文件名加上序号和标题名命名，每次遇到新标题序号 +1 （注意：文件名长度可能受到操作系统的限制，导致读取失败，有相关问题修改下标题名）
            current_title = paragraph.text
            new_doc = Document()
            new_doc.add_paragraph(paragraph.text, style=paragraph.style.name)
        elif new_doc is not None:
            new_paragraph = new_doc.add_paragraph(paragraph.text, style='Normal')

            # 遍历段落中的运行元素
            for run in paragraph.runs:
                # 获取运行元素的XML表示
                run_xml = run._r
                # 检查是否包含图片
                if '<w:drawing>' in run_xml.xml:
                    print("found an image1")
                    # 解析图片信息
                    image_start = run_xml.xml.find('<w:drawing>')
                    image_end = run_xml.xml.find('</w:drawing>') + len('</w:drawing>')
                    image_xml = run_xml.xml[image_start:image_end]

                    # 获取关联标识符
                    image_id_start = image_xml.find('r:embed="') + len('r:embed="')
                    image_id_end = image_xml.find('"', image_id_start)
                    image_id = image_xml[image_id_start:image_id_end]
                    # print(image_id)
                    image_path = 'images/' + image_id + '.png'
                    print("image_path=" + image_path)

                    # 根据image_id，将images文件夹的图片复制到新文档中
                    new_paragraph.add_run().add_picture(image_path, width=Inches(5))

                if not new_paragraph.runs:  # 如果段落中没有图片，则添加原文本内容
                    new_paragraph.text = paragraph.text

    # 保存新文档
    if new_doc is not None:
        file_name = current_title + '.docx'
        file_path = os.path.join(file_folder, file_name)
        new_doc.save(file_path)


# 使用示例
#extract_and_save_images_from_docx()  # 替换你要分化的文档
extract_images_from_docx('C:\\Users\\10109\\Desktop\\横向_政策抽取\\25350-关于开展2022年第二批云浮市创新创业团队和创新创业人才项目申报的通知.docx')  # 替换你要分化的文档