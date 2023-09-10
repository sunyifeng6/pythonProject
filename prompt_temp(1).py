from langchain.prompts.few_shot import FewShotPromptTemplate
from langchain.prompts.prompt import PromptTemplate
# from langchain.llms import OpenAI
from langchain.chat_models import ChatOpenAI
import os
from docx import Document

import docx


def getText(fileName):
    doc = docx.Document(fileName)
    TextList = []
    for paragraph in doc.paragraphs:
        TextList.append(paragraph.text)

    return '\n'.join(TextList)

os.environ["OPENAI_API_KEY"] = "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"
# 有两个示例
examples = [

]

example_prompt = PromptTemplate(input_variables=["question", "answer"], template="Question: {question}\n{answer}")

prompt = FewShotPromptTemplate(
    examples=examples,
    example_prompt=example_prompt,
    prefix="你是一个文档切分器",
    suffix="Question: 请你从以下文本中利用{project},每个文档格式必须要与原文格式保持一致，每个文档的标题数需与原文保持一致，每个文档都不能遗漏原文的信息，并且文档以'end'结束，用来区分不同文档:{input}",
    input_variables=["project","input"]
)
# 构建模型
chat = ChatOpenAI(temperature=0,model="gpt-3.5-turbo")

# project="创新创业团队、创新人才、创业人才三个项目对应生成三个文档"
# input="八、联系方式联系人及电话：\n成果专家科 张舒怡，8921136"

project="创新创业团队、创新人才、创业人才三个项目对应生成三个文档"
# input="三、申报条件\n（一)创新创业团队创新创业团队是指依托产业化项目引进，引进后在引进企业牵头实施创新创业项目，重点围绕我市“园区经济、镇域经济、资源经济”（下称“三大经济”）产业的创新技术产业化发展，掌握产业领域关键核心技术或具有知识产权，拥有技术成熟的成果并能进入开发阶段，或引进后研发有效解决新产品、新工艺、新技术难题。具备切实可行的未来发展计划，产业发展潜力和市场化空间较大的创新创业团队。必须满足以下全部条件：\n1.团队应包括1名负责人和不少于2名核心成员（指与团队带头人同时引进的人才，下同），团队人数不超过6名。团队成员应具有硕士研究生以上学历学位或副高以上专业技术职称。年龄原则上不超过55周岁。\n2.团队成员具有合理的专业结构，团队负责人或至少1名核心成员引进前应在国内外知名高校、科研院所工作且有相当于正高专业技术职称，或是在国内外知名企业担任中高级技术管理领导职务2 年以上，且前期应在“三大经济”相关项目、产品和技术开发等方面至少有1年以上较好的研究成果和成果转化业绩。\n3.团队掌握的核心技术应具有行业一流水平，技术成熟并已进入产业化阶段，能较大程度推动我市“三大经济”发展。团队成员在技术、应用、产业化等环节有关联性，具备达到创新创业目标的完整创新链。\n4.团队引进企业应当符合我市“三大经济”发展方向。引进后创新创业团队须承诺在云浮工作时间不少于3年，其中1名核心成员全职在云浮工作，其他团队成员在云浮工作时间每年累计不少于2个月。\n5.团队引进企业经营运行状况良好，且技术创新体系完善，研发组织完备，企业研发费用占主营业务收入比重一般应高于3%，为团队从事的项目配足科研资金，提供相关研发设备，落实项目产业化所需的各类要素，为团队成员提供良好的生活保障。团队引进企业还应具备我市“三大经济”研究实践基础，或在攻关的技术方向上拥有国际、国内领先技术成果，具备创新创业的基础设施和条件。\n6.创新创业团队是关联企业派出的，引进时必须转移创新技术成果在云浮应用或带动新产品在云浮生产，且近两年或引进后承诺2年内该技术成果应用投入或产值不少于5000万元，派驻期间研究成果归属引进企业。\n（二）创新人才创新人才是指我市国有企业和民营企业引进的市外优秀专业技术技能人才，应当取得博士学位或具有副高以上专业技术职称，年龄原则上不超过55周岁。要求全职引进后承诺在云浮工作时间不少 于3年。同时符合下列条件之一：\n1.曾在国内外综合排名前500强企业、行业排名前500强企业担任高级职务的专业技术人才；已与市内企业签约或意向签约，或已成为市内企业技术、商业主要合作伙伴的人才。\n2.在市外高校、科研机构担任相当于副教授及以上职务、具有较强科技创新能力的专家学者。\n3.掌握产业领域关键核心技术或能解决产品生产工艺难题的人才。\n（三）创业人才创业人才是指掌握领先技术的来云创业人员。应取得全日制本科以上学历学位，年龄原则上不超过55周岁，个人投资额不少于200万元，企业须完成企业工商登记，主要创办人为第一大股东或者最 大自然人股东。创办的企业应符合我市金属智造、生物医药、现代农业等十大产业发展方向，投资额1000万元以上，产品具有核心技术可直接进行产业化，或已完成前期开发，进入中试或样品制造阶 段，具有稳定实验数据和中试产品的项目。引进后承诺在云浮工作时间不少于3年，不硬性要求全职 引进。同时符合下列条件之一：\n1.创业人才拥有国际或国内领先技术成果。\n2.创业人才有相关创业经验或者曾任国内外知名高校、科研院所担任中高层管理职务2年以上， 有较强的经营管理能力，具有主持研发成果成功实施产业化的经历。\n3.拥有广阔市场前景的自主创新产品，获得风险投资机构投资。"
input = getText("零散-25350/二、专题类别.docx")


_input = prompt.format(project=project,input=input)
# print(_input)
output = chat.predict(_input)
sections = output.split("end\n")
for i, section in enumerate(sections):
    doc = Document()
    doc.add_paragraph(section.strip())
    doc.save(f"C:\\Users\\10109\\Desktop\\横向_政策抽取\\out\\document_{i+1}.docx")